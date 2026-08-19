#!/usr/bin/env python3
"""Style the pandoc-generated walkthrough.docx: cover page, heading colors,
shaded code blocks, table header row, page numbers in the footer, and
red-colored command lines.

Shell commands (```bash fences) render in dark red so a reader can pick out
"type this" lines at a glance; file-content blocks (json/html/js/toml) and
command *output* (indented code blocks) stay black. Pandoc merges and wraps
code lines into paragraphs unpredictably, so each paragraph is matched back
to its md code block by text (a paragraph never spans two blocks); ```bash
blocks mark commands.

Usage: pandoc walkthrough.md -o walkthrough.docx --toc --toc-depth=2
       python3 scripts/style-walkthrough-docx.py docs/walkthrough.docx
"""
import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HEADING_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # dark professional blue
CODE_FILL = "F2F2F2"                         # light gray behind code blocks
TABLE_HEADER_FILL = "DEEAF6"                 # light blue table header row
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
COMMAND_RED = RGBColor(0xC0, 0x00, 0x00)     # dark red for command lines


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        run.font.size = Pt(9)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), r"PAGE \* MERGEFORMAT")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        p._p.append(fld)
        run2 = p.add_run(" of ")
        run2.font.size = Pt(9)
        fld2 = OxmlElement("w:fldSimple")
        fld2.set(qn("w:instr"), r"NUMPAGES \* MERGEFORMAT")
        r2 = OxmlElement("w:r")
        t2 = OxmlElement("w:t")
        t2.text = "1"
        r2.append(t2)
        fld2.append(r2)
        p._p.append(fld2)


def add_cover_page(doc):
    first = doc.paragraphs[0]
    spacer = first.insert_paragraph_before("")
    spacer.add_run("").add_break(WD_BREAK.LINE)  # breathing room
    title = first.insert_paragraph_before("DocuTrust")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("DocuTrust")
    tr.font.size = Pt(40)
    tr.font.bold = True
    tr.font.color.rgb = HEADING_COLOR

    subtitle = first.insert_paragraph_before("DevSecOps Project 1 — Walkthrough")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("DevSecOps Project 1 — Walkthrough")
    sr.font.size = Pt(20)
    sr.font.color.rgb = ACCENT

    tagline = first.insert_paragraph_before(
        "SAST · Secrets Scanning · Live Secret Verification"
    )
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tg = tagline.add_run("SAST · Secrets Scanning · Live Secret Verification")
    tg.font.size = Pt(13)
    tg.font.italic = True

    rule = first.insert_paragraph_before("")
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = rule.add_run("— " * 12)
    rl.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)

    note = first.insert_paragraph_before(
        "Every command in this walkthrough actually ran, and every output "
        "quoted is real captured output from the tools — with evidence files "
        "and screenshots cited along the way."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Every command in this walkthrough actually ran, and every output "
        "quoted is real captured output from the tools — with evidence files "
        "and screenshots cited along the way."
    )
    nr.font.size = Pt(10)
    nr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    date_p = first.insert_paragraph_before("August 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run("August 2026")
    dr.font.size = Pt(11)

    # page break after the cover
    brk = first.insert_paragraph_before("")
    brk.add_run("").add_break(WD_BREAK.PAGE)


def find_style(doc, name):
    # pandoc's reference docx trips python-docx's indexed lookup; iterate
    for st in doc.styles:
        if st.name == name:
            return st
    return None


def parse_md_code(md_path):
    """Parse walkthrough.md into its code blocks: fenced blocks (with their
    language tag) and 4-space-indented blocks (pandoc's indented code, used
    here for command *output*). Commands are fenced ```bash blocks."""
    blocks = []  # each: {"lang", "cmd": bool, "lines": [str, ...]}
    in_fence = False
    fence_lang = None
    cur = None
    with open(md_path, encoding="utf-8") as f:
        for raw in f:
            s = raw.rstrip("\n")
            if s.strip().startswith("```"):
                if not in_fence:
                    in_fence = True
                    fence_lang = s.strip()[3:].strip()  # "" for bare fences
                    cur = []
                else:
                    blocks.append({"lang": fence_lang,
                                   "cmd": fence_lang == "bash", "lines": cur})
                    in_fence, cur = False, None
                continue
            if in_fence:
                cur.append(s)  # fenced: keep indentation verbatim
            elif cur is None and len(s) - len(s.lstrip()) >= 4 and s.strip():
                # 4-space-indented block: pandoc strips the common indent
                cur = [s]
            elif cur is not None and (s.strip() == "" or
                                      len(s) - len(s.lstrip()) >= 4):
                cur.append(s)
            else:
                if cur is not None:
                    indent = min((len(l) - len(l.lstrip())
                                  for l in cur if l.strip()), default=0)
                    blocks.append({"lang": None, "cmd": False,
                                   "lines": [l[indent:] for l in cur]})
                    cur = None
    if cur is not None and not in_fence:
        indent = min((len(l) - len(l.lstrip()) for l in cur if l.strip()),
                     default=0)
        blocks.append({"lang": None, "cmd": False,
                       "lines": [l[indent:] for l in cur]})
    return blocks


def match_paragraphs(paras, blocks):
    """Map each docx Source Code paragraph back to the md block it came from.

    Pandoc merges lines into one paragraph (w:br breaks) and wraps long lines,
    so line counts don't line up — but a paragraph never spans two code
    blocks. Match each paragraph's text segments greedily against the md
    lines in order; the block's cmd flag decides the color."""
    b = l = c = 0            # pointer: block b, line l, char c
    out = []
    for t in paras:
        segs = t.split("\n")
        target_b = None
        for si in segs:
            if si == "":
                continue  # blank line: consumed at current position
            hit = None
            while b < len(blocks) and hit is None:
                lines = blocks[b]["lines"]
                if l >= len(lines):
                    b, l, c = b + 1, 0, 0
                    continue
                rest = lines[l][c:]
                if rest.startswith(si):
                    hit = (b, l, c)
                else:
                    # maybe the segment begins on a later line of this block
                    for j in range(l + 1, len(lines)):
                        if lines[j].startswith(si):
                            hit = (b, j, 0)
                            break
                    if hit is None:
                        b, l, c = b + 1, 0, 0
            if hit is None:
                raise SystemExit(f"could not match code paragraph: {t!r}")
            hb, hl, hc = hit
            if target_b is None:
                target_b = hb
            # advance past the matched segment
            l, c = hl, hc + len(si)
            if c >= len(blocks[hb]["lines"][l]):
                l, c = l + 1, 0
        if target_b is None:
            target_b = b     # blank paragraph: belongs to the current block
        out.append(blocks[target_b]["cmd"])
    # fully consumed trailing blocks don't count as leftover
    while b < len(blocks) and l >= len(blocks[b]["lines"]):
        b += 1
        l = 0
    if b < len(blocks):
        raise SystemExit(f"leftover md code after {len(paras)} paragraphs "
                         f"(block {b} of {len(blocks)})")
    return out


def main(path):
    doc = Document(path)

    # --- styles ---------------------------------------------------------
    for name, size in (("Heading 1", 17), ("Heading 2", 13)):
        st = find_style(doc, name)
        if st is None:
            continue
        st.font.color.rgb = HEADING_COLOR
        st.font.size = Pt(size)
        st.font.bold = True

    # --- code blocks: shaded, tighter; commands in red -----------------
    # Pandoc's paragraph layout doesn't line up 1:1 with md lines, so match
    # each paragraph back to its md block; ```bash blocks are commands.
    md_path = os.path.join(os.path.dirname(os.path.abspath(path)),
                           "walkthrough.md")
    blocks = parse_md_code(md_path)
    paras = [p for p in doc.paragraphs
             if p.style and ("Source Code" in p.style.name
                             or "SourceCode" in p.style.name)]
    marks = match_paragraphs([p.text for p in paras], blocks)
    for p, is_cmd in zip(paras, marks):
        shade_paragraph(p, CODE_FILL)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if is_cmd:
            for r in p.runs:
                r.font.color.rgb = COMMAND_RED

    # --- tables: header row shaded + bold -----------------------------
    for tbl in doc.tables:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in tbl.rows[0].cells:
            shade_cell(cell, TABLE_HEADER_FILL)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True

    add_page_number_footer(doc)
    add_cover_page(doc)
    doc.save(path)
    print(f"styled {path}: cover, heading colors, code shading (commands in red), "
          "table headers, footer page numbers")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "docs/walkthrough.docx")
