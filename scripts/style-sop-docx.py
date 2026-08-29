#!/usr/bin/env python3
"""Style a pandoc-generated SOP docx to the DocuTrust professional standard:
per-project cover page (SOP variant with version/author/date), table of
contents, colored headings, accent-bordered code blocks (commands in dark
red), callout boxes for quotes, styled figure captions, bordered tables
with shaded headers, and a footer with the document title + page numbers.

Sibling of style-walkthrough-docx.py: same visual language, SOP cover.

Usage (from the SOP's output directory, so relative image paths resolve):
    pandoc Project-1-SOP.md -o Project-1-SOP.docx --toc --toc-depth=2
    python3 ../../../scripts/style-sop-docx.py Project-1-SOP.docx [source.md]
"""
import glob
import os
import re
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
GOLD = RGBColor(0x8A, 0x66, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)
RULER = RGBColor(0xC9, 0xC9, 0xC9)
LIGHT = "F2F4F7"
CALL = "FFF8E7"
CALL_BAR = "E0B93C"
HEADER_FILL = "DEEAF6"
ZEBRA = "F7F9FC"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

COVER_TRACKS = {
    1: ("Project 1 — SAST & Secrets Scanning",
        "SAST · Secrets Scanning · Live Secret Verification"),
    2: ("Project 2 — SCA & Dependency Risk",
        "SCA · Dependency Confusion · OpenSSF Scorecard"),
    3: ("Project 3 — Runtime Testing & Defense",
        "DAST · IAST · RASP — Compared for Real"),
}


def shade_paragraph(p, fill, accent=None):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    if accent is not None:
        pbdr = OxmlElement("w:pBdr")
        el = OxmlElement("w:left")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "24")
        el.set(qn("w:space"), "10")
        el.set(qn("w:color"), accent)
        pbdr.append(el)
        pPr.append(pbdr)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_table_borders(tbl, color="BFBFBF", sz="6"):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def add_footer(doc, doc_title):
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def rn(text=None, size=8.5, color=GREY):
            r = p.add_run(text or "")
            r.font.size = Pt(size)
            r.font.color.rgb = color
            r.font.name = BODY_FONT
            return r

        rn(doc_title + "   —   ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), r"PAGE \* MERGEFORMAT")
        rr = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        rr.append(t)
        fld.append(rr)
        p._p.append(fld)
        rn(" of ")
        fld2 = OxmlElement("w:fldSimple")
        fld2.set(qn("w:instr"), r"NUMPAGES \* MERGEFORMAT")
        rr2 = OxmlElement("w:r")
        t2 = OxmlElement("w:t")
        t2.text = "1"
        rr2.append(t2)
        fld2.append(rr2)
        p._p.append(fld2)


def add_cover_page(doc, project_no):
    anchor = None
    for p in doc.paragraphs:
        if p.style and "Title" in (p.style.name or ""):
            anchor = p
            break
    if anchor is None:
        anchor = doc.paragraphs[0]
    toc = None
    for p in doc.paragraphs:
        if "Table of Contents" in p.text:
            toc = p
            break

    made = []

    def line(text="", size=10, color=None, bold=False, italic=False,
             page_break=False):
        p = anchor.insert_paragraph_before("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = BODY_FONT
        if color is not None:
            run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic
        if page_break:
            p.add_run("").add_break(WD_BREAK.PAGE)
        made.append(p._p)
        return p

    line("DOCUTRUST — DEVSECOPS TRACK · CHAIN A", 9, GOLD, bold=True)
    line("", 13)
    line("DocuTrust", 42, NAVY, bold=True)
    line(COVER_TRACKS[project_no][0], 19, BLUE)
    line(COVER_TRACKS[project_no][1], 12, GREY, italic=True)
    line("— " * 14, 12, RULER)
    line("DevSecOps Implementation Guide / SOP", 14, NAVY, bold=True)
    line("Version 1.0 · Author: Codelak · August 2026", 11, GREY)
    line("", 10)
    line(
        "Every command in this guide actually ran, and every figure is a "
        "real capture of that command on a real terminal. An engineer can "
        "execute this guide from start to finish — environment preparation, "
        "task-by-task implementation, verification, CI gates, and evidence — "
        "and reach the same final state.",
        10, GREY, italic=True)
    line(
        "Format: task-based · every command shown with its expected output · "
        "checkpoints before every task · rerun and idempotency documented.",
        9.5, BLUE)
    line("DocuTrust track · updated August 2026", 10, GREY)
    line("", 11, page_break=True)

    # The cover must be the very first thing. Pandoc writes the TOC as a
    # structured-document tag (sdt) at the top of the body, so lift the
    # cover paragraphs above it; the TOC field stays intact on page 2.
    sdt = doc.element.body.find(qn("w:sdt"))
    if sdt is not None:
        for el in made:
            sdt.addprevious(el)


def find_style(doc, name):
    for st in doc.styles:
        if st.name == name:
            return st
    return None


def parse_md_code(md_path):
    blocks = []
    in_fence = False
    fence_lang = None
    cur = None
    with open(md_path, encoding="utf-8") as f:
        for raw in f:
            s = raw.rstrip("\n")
            if s.strip().startswith("```"):
                if not in_fence:
                    in_fence = True
                    fence_lang = s.strip()[3:].strip()
                    cur = []
                else:
                    blocks.append({"lang": fence_lang,
                                   "cmd": fence_lang == "bash", "lines": cur})
                    in_fence, cur = False, None
                continue
            if in_fence:
                cur.append(s)
            elif cur is None and len(s) - len(s.lstrip()) >= 4 and s.strip():
                cur = [s]
            elif cur is not None and (s.strip() == "" or
                                      len(s) - len(s.lstrip()) >= 4):
                cur.append(s)
            else:
                if cur is not None:
                    indent = min((len(l) - len(l.lstrip())
                                  for l in cur if l.strip()), default=0)
                    blocks.append({"lang": None, "cmd": False,
                                   "lines": [l[indent:] for l in cur]})
                    cur = None
    if cur is not None and not in_fence:
        indent = min((len(l) - len(l.lstrip()) for l in cur if l.strip()),
                     default=0)
        blocks.append({"lang": None, "cmd": False,
                       "lines": [l[indent:] for l in cur]})
    return blocks


def match_paragraphs(paras, blocks):
    b = l = c = 0
    out = []
    for t in paras:
        segs = t.split("\n")
        target_b = None
        for si in segs:
            if si == "":
                continue
            hit = None
            while b < len(blocks) and hit is None:
                lines = blocks[b]["lines"]
                if l >= len(lines):
                    b, l, c = b + 1, 0, 0
                    continue
                rest = lines[l][c:]
                if rest.startswith(si):
                    hit = (b, l, c)
                else:
                    for j in range(l + 1, len(lines)):
                        if lines[j].startswith(si):
                            hit = (b, j, 0)
                            break
                    if hit is None:
                        b, l, c = b + 1, 0, 0
            if hit is None:
                raise SystemExit(f"could not match code paragraph: {t!r}")
            hb, hl, hc = hit
            if target_b is None:
                target_b = hb
            l, c = hl, hc + len(si)
            if c >= len(blocks[hb]["lines"][l]):
                l, c = l + 1, 0
        if target_b is None:
            target_b = b
        out.append(blocks[target_b]["cmd"])
    while b < len(blocks) and l >= len(blocks[b]["lines"]):
        b += 1
        l = 0
    if b < len(blocks):
        raise SystemExit(
            f"leftover md code after {len(paras)} paragraphs (block {b} of "
            f"{len(blocks)})")
    return out


def set_font(p, name=BODY_FONT, size=None, color=None, bold=None):
    for r in p.runs:
        r.font.name = name
        if size is not None:
            r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        if bold is not None:
            r.font.bold = bold


def main(path, md_path=None):
    doc = Document(path)
    if md_path is None:
        base = os.path.dirname(os.path.abspath(path))
        candidates = sorted(glob.glob(os.path.join(base, "*.md")))
        md_path = candidates[0]
    title = ""
    with open(md_path, encoding="utf-8") as f:
        for line in f:
            if line.startswith("# "):
                title = line[2:].strip()
                break
    project_no = 1
    m = re.search(r"Project (\d)", title)
    if m:
        project_no = int(m.group(1))

    for name, size in (("Heading 1", 16.5), ("Heading 2", 13.5),
                       ("Heading 3", 11.5), ("Title", 22)):
        st = find_style(doc, name)
        if st is None:
            continue
        st.font.color.rgb = NAVY
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.name = BODY_FONT

    body = find_style(doc, "Body Text") or find_style(doc, "Normal")
    if body is not None:
        body.font.name = BODY_FONT
        body.font.size = Pt(10.5)

    blocks = parse_md_code(md_path)
    paras = [p for p in doc.paragraphs
             if p.style and ("Source Code" in p.style.name
                             or "SourceCode" in p.style.name)]
    marks = match_paragraphs([p.text for p in paras], blocks)
    for p, is_cmd in zip(paras, marks):
        shade_paragraph(p, LIGHT)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.right_indent = Cm(0.15)
        set_font(p, MONO_FONT, 9.5)
        if is_cmd:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    for p in doc.paragraphs:
        if p.style and ("Quote" in p.style.name
                        or "BlockText" in (p.style.style_id or "")):
            shade_paragraph(p, CALL, accent=CALL_BAR)
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.right_indent = Cm(0.25)
            set_font(p, BODY_FONT, 10)
        elif p.style and "Caption" in (p.style.name or ""):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, BODY_FONT, 9, GREY)
            for r in p.runs:
                r.font.italic = True
        elif p._p.findall(".//" + qn("w:drawing")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for tbl in doc.tables:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl)
        for cell in tbl.rows[0].cells:
            shade_cell(cell, HEADER_FILL)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        for r_i, row in enumerate(tbl.rows[1:], start=1):
            if r_i % 2 == 0:
                for cell in row.cells:
                    shade_cell(cell, ZEBRA)

    add_footer(doc, f"DocuTrust — {title}")
    add_cover_page(doc, project_no)
    doc.save(path)
    print(f"styled {path}: SOP cover, TOC, headings, code boxes, callouts, "
          "captions, tables, footer")


if __name__ == "__main__":
    args = sys.argv[1:]
    main(args[0], args[1] if len(args) > 1 else None)
